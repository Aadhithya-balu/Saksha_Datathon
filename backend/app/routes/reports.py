"""Live reporting routes with paginated data, PDF export, DOCX export, TXT export, and CSV export.

Issue #176 adds a production audit lifecycle on top of the legacy live exports:
draft -> generate -> review -> finalize -> archive with source/evidence linking,
provenance, versioning, integrity hashing and immutable finalized snapshots.
"""
from __future__ import annotations

import csv
import io
import json
import uuid
from datetime import datetime, timezone
from typing import Any, Callable
from uuid import UUID

from fpdf import FPDF
from docx import Document
from pydantic import BaseModel

from fastapi import APIRouter, Depends, Query, Request
from fastapi.responses import Response
from sqlalchemy import asc, desc, func, or_
from sqlalchemy.orm import Query as SQLAlchemyQuery
from sqlalchemy.orm import Session, joinedload

from app.auth.dependencies import get_current_user
from app.auth.rbac import ROLE_ADMIN, ROLE_CRIME_ANALYST, ROLE_INSPECTOR, ROLE_INVESTIGATOR, ROLE_POLICYMAKER, require_roles
from app.core.exceptions import ConflictException, NotFoundException
from app.database.postgres import get_db
from app.models.audit_log import AuditLog
from app.models.crime import CrimeCase
from app.models.criminal import Criminal
from app.models.evidence import Evidence
from app.models.location import Location
from app.models.officer import Officer
from app.models.report import (
    GEN_METHOD_DATABASE_EXPORT,
    REPORT_STATUS_GENERATED,
    Report,
)
from app.models.user import User
from app.schemas.common import PaginatedResponse
from app.schemas.report import (
    ReportCreateRequest,
    ReportDetailOut,
    ReportGeneratePayload,
    ReportOut,
    ReportReviewRequest,
    ReportValidateRequest,
    ReportVersionCreateRequest,
)
from app.services import audit_service, report_service

router = APIRouter(prefix="/reports", tags=["Reports"], dependencies=[Depends(require_roles(
    ROLE_ADMIN,
    ROLE_CRIME_ANALYST,
    ROLE_INVESTIGATOR,
    ROLE_INSPECTOR,
    ROLE_POLICYMAKER,
))])

REPORT_TYPES = {"cases", "officers", "criminals", "evidence"}
EXPORT_FORMATS = {"pdf", "csv", "docx", "txt", "xlsx"}
SORTABLE_COLUMNS: dict[str, dict[str, Any]] = {
    "cases": {
        "case_number": CrimeCase.case_number,
        "occurred_at": CrimeCase.occurred_at,
        "reported_at": CrimeCase.reported_at,
        "status": CrimeCase.status,
        "priority": CrimeCase.priority,
        "created_at": CrimeCase.created_at,
    },
    "officers": {
        "name": Officer.name,
        "badge_number": Officer.badge_number,
        "district": Officer.district,
        "station": Officer.station,
        "status": Officer.status,
        "created_at": Officer.created_at,
    },
    "criminals": {
        "full_name": Criminal.full_name,
        "status": Criminal.status,
        "gender": Criminal.gender,
        "created_at": Criminal.created_at,
    },
    "evidence": {
        "title": Evidence.title,
        "evidence_type": Evidence.evidence_type,
        "status": Evidence.status,
        "created_at": Evidence.created_at,
    },
}


def _format_human_readable_value(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, (int, float, bool)):
        return str(v)
    if isinstance(v, str):
        return v.strip()
    if isinstance(v, dict):
        if v.get("type") == "Feature" and "properties" in v:
            props = v.get("properties", {})
            coords = v.get("geometry", {}).get("coordinates", [])
            coord_str = f" (Coords: {coords[0]}, {coords[1]})" if coords and len(coords) >= 2 else ""
            items = [f"{str(pk).replace('_', ' ').title()}: {pv}" for pk, pv in props.items()]
            return ", ".join(items) + coord_str
        res = []
        for k, val in v.items():
            key_name = str(k).replace("_", " ").title()
            res.append(f"- {key_name}: {_format_human_readable_value(val)}")
        return "\n".join(res)
    if isinstance(v, list):
        res = []
        for idx, item in enumerate(v, 1):
            if isinstance(item, dict):
                res.append(f"{idx}. {_format_human_readable_value(item)}")
            else:
                res.append(f"- {str(item)}")
        return "\n".join(res)
    return str(v)


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return _format_human_readable_value(value)
    return str(value).strip()


def _serialize_datetime(value: Any) -> str:
    if isinstance(value, datetime):
        return value.astimezone(timezone.utc).isoformat()
    return _clean_text(value)


def _apply_date_range(query: SQLAlchemyQuery, column: Any, date_from: datetime | None, date_to: datetime | None):
    if date_from:
        query = query.filter(column >= date_from)
    if date_to:
        query = query.filter(column <= date_to)
    return query


def _apply_sort(query: SQLAlchemyQuery, report_type: str, sort_by: str, sort_order: str):
    column = SORTABLE_COLUMNS[report_type].get(sort_by)
    if column is None:
        column = SORTABLE_COLUMNS[report_type]["created_at"]
    direction = asc if sort_order == "asc" else desc
    return query.order_by(direction(column))


def _report_query(
    db: Session,
    report_type: str,
    search: str | None,
    status: str | None,
    district: str | None,
    date_from: datetime | None,
    date_to: datetime | None,
    sort_by: str,
    sort_order: str,
) -> tuple[SQLAlchemyQuery, list[str], Callable[[Any], dict[str, Any]]]:
    if report_type == "cases":
        query = db.query(CrimeCase).options(
            joinedload(CrimeCase.assigned_officer),
            joinedload(CrimeCase.category),
            joinedload(CrimeCase.location)
        )
        if search:
            query = query.filter(or_(CrimeCase.case_number.ilike(f"%{search}%"), CrimeCase.description.ilike(f"%{search}%")))
        if status:
            query = query.filter(CrimeCase.status == status)
        if district:
            query = query.join(Location, CrimeCase.location_id == Location.id, isouter=True).filter(Location.district == district)
        query = _apply_date_range(query, CrimeCase.occurred_at, date_from, date_to)
        headers = ["case_number", "category", "district", "station", "status", "priority", "progress", "occurred_at", "reported_at", "assigned_officer", "mo_tags", "description"]
        def mapper(item):
            return {
            "case_number": item.case_number,
            "category": item.category.name if item.category else "",
            "district": item.location.district if item.location else "",
            "station": item.location.station if item.location else "",
            "status": item.status,
            "priority": item.priority or "medium",
            "progress": f"{item.progress or 0}%",
            "occurred_at": _serialize_datetime(item.occurred_at),
            "reported_at": _serialize_datetime(item.reported_at),
            "assigned_officer": item.assigned_officer.name if item.assigned_officer else "Unassigned",
            "mo_tags": item.mo_tags or "",
            "description": item.description or "",
        }
    elif report_type == "officers":
        query = db.query(Officer)
        if search:
            query = query.filter(or_(Officer.name.ilike(f"%{search}%"), Officer.badge_number.ilike(f"%{search}%"), Officer.email.ilike(f"%{search}%")))
        if status:
            query = query.filter(Officer.status == status)
        if district:
            query = query.filter(Officer.district == district)
        headers = ["badge_number", "name", "rank", "designation", "district", "station", "status", "phone", "email"]
        def mapper(item):
            return {
            "badge_number": item.badge_number,
            "name": item.name,
            "rank": item.rank or "",
            "designation": item.designation or "",
            "district": item.district or "",
            "station": item.station or "",
            "status": item.status,
            "phone": item.phone or "",
            "email": item.email or "",
        }
    elif report_type == "criminals":
        query = db.query(Criminal)
        if search:
            query = query.filter(or_(Criminal.full_name.ilike(f"%{search}%"), Criminal.aliases.ilike(f"%{search}%"), Criminal.mo_summary.ilike(f"%{search}%")))
        if status:
            query = query.filter(Criminal.status == status)
        headers = ["full_name", "aliases", "gender", "date_of_birth", "status", "address", "identifying_marks", "mo_summary"]
        def mapper(item):
            return {
            "full_name": item.full_name,
            "aliases": item.aliases or "",
            "gender": item.gender or "",
            "date_of_birth": _serialize_datetime(item.date_of_birth),
            "status": item.status,
            "address": item.address or "",
            "identifying_marks": item.identifying_marks or "",
            "mo_summary": item.mo_summary or "",
        }
    else:
        query = db.query(Evidence).options(joinedload(Evidence.crime_case), joinedload(Evidence.assignee))
        if search:
            query = query.filter(or_(Evidence.title.ilike(f"%{search}%"), Evidence.description.ilike(f"%{search}%"), Evidence.evidence_type.ilike(f"%{search}%")))
        if status:
            query = query.filter(Evidence.status == status)
        query = _apply_date_range(query, Evidence.created_at, date_from, date_to)
        headers = ["title", "case_number", "evidence_type", "status", "assigned_to", "created_by", "storage_path", "created_at", "description"]
        def mapper(item):
            return {
            "title": item.title,
            "case_number": item.crime_case.case_number if item.crime_case else "",
            "evidence_type": item.evidence_type,
            "status": item.status,
            "assigned_to": item.assignee.full_name if item.assignee else "",
            "created_by": item.created_by or "",
            "storage_path": item.storage_path or "",
            "created_at": _serialize_datetime(item.created_at),
            "description": item.description or "",
        }
    return _apply_sort(query, report_type, sort_by, sort_order), headers, mapper


def _filters_dict(**filters: Any) -> dict[str, Any]:
    return {key: _serialize_datetime(value) for key, value in filters.items() if value not in (None, "")}


def _csv_response(filename: str, headers: list[str], rows: list[dict[str, Any]]) -> Response:
    buffer = io.StringIO()
    buffer.write("\ufeff")
    writer = csv.DictWriter(buffer, fieldnames=headers, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow({header: _clean_text(row.get(header)) for header in headers})
    return Response(
        content=buffer.getvalue().encode("utf-8"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}.csv"'},
    )


def _xlsx_response(filename: str, title: str, filters: dict, headers: list[str], rows: list[dict[str, Any]]) -> Response:
    """Native Excel workbook export (gap M1 — previously unimplemented)."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"

    header_fill = PatternFill(start_color="FF1E293B", end_color="FF1E293B", fill_type="solid")
    header_font = Font(color="FFF8FAFC", bold=True, size=10)

    # Metadata block
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    filter_str = ", ".join(f"{k}={v}" for k, v in filters.items()) if filters else "None"
    sheet.append(["SAKSHA Police Intelligence & Analytics Platform"])
    sheet.append([title])
    sheet.append([f"Generated At: {generated} | Filters: {filter_str} | Total Records: {len(rows)}"])
    sheet.append([])
    meta_rows = 4

    formatted_headers = [h.replace("_", " ").title() for h in headers]
    sheet.append(formatted_headers)
    for cell in sheet[meta_rows]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(vertical="center")

    for row in rows:
        sheet.append([_clean_text(row.get(h))[:400] for h in headers])

    for index, header in enumerate(headers, start=1):
        max_len = min(max((len(str(sheet.cell(row=r, column=index).value or "")) for r in range(meta_rows, min(sheet.max_row, meta_rows + 200) + 1)), default=10) + 2, 60)
        sheet.column_dimensions[get_column_letter(index)].width = max(12, max_len)
    sheet.freeze_panes = f"A{meta_rows + 1}"

    buffer = io.BytesIO()
    workbook.save(buffer)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}.xlsx"'},
    )


def _generate_txt(title: str, filters: dict, headers: list[str], rows: list[dict]) -> bytes:
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    filter_str = ", ".join(f"{k}={v}" for k, v in filters.items()) if filters else "None"
    
    lines = [
        "=" * 90,
        "SAKSHA POLICE INTELLIGENCE & ANALYTICS PLATFORM".center(90),
        "CONFIDENTIAL LAW-ENFORCEMENT REPORT".center(90),
        "=" * 90,
        "",
        f"REPORT TITLE   : {title.upper()}",
        f"GENERATED AT   : {generated}",
        f"TOTAL RECORDS  : {len(rows)}",
        f"FILTERS APPLIED: {filter_str}",
        "",
        "-" * 90,
        "REPORT SUMMARY & DETAILS".center(90),
        "-" * 90,
        ""
    ]
    
    if not headers or not rows:
        lines.append("No records available for the current filter criteria.")
    else:
        formatted_headers = [h.replace("_", " ").title() for h in headers]
        col_widths = {h: len(fh) for h, fh in zip(headers, formatted_headers)}
        for row in rows:
            for h in headers:
                val = _clean_text(row.get(h))
                col_widths[h] = max(col_widths[h], len(val))
        
        for h in headers:
            col_widths[h] = min(max(col_widths[h], 10), 35)
            
        header_line = " | ".join(fh.ljust(col_widths[h]) for h, fh in zip(headers, formatted_headers))
        lines.append(header_line)
        lines.append("-" * len(header_line))
        
        for row in rows:
            line_parts = []
            for h in headers:
                val = _clean_text(row.get(h)).replace('\n', ' ')
                if len(val) > col_widths[h]:
                    val = val[:col_widths[h] - 3] + "..."
                line_parts.append(val.ljust(col_widths[h]))
            lines.append(" | ".join(line_parts))
            
    lines.append("")
    lines.append("=" * 90)
    lines.append("SECURITY COMPLIANCE ACT NOTICE: CONFIDENTIAL LAW-ENFORCEMENT REPORT".center(90))
    lines.append("=" * 90)
    
    return "\n".join(lines).encode("utf-8")


def _generate_pdf(title: str, filters: dict, headers: list[str], rows: list[dict]) -> bytes:
    import os as _os
    import unicodedata

    # Typography replacements for PDF encoding safety
    UNICODE_REPLACEMENTS = {
        "\u2014": " - ",   # em-dash
        "\u2013": " - ",   # en-dash
        "\u2022": " * ",   # bullet
        "\u2018": "'",     # left single quote
        "\u2019": "'",     # right single quote
        "\u201c": '"',     # left double quote
        "\u201d": '"',     # right double quote
        "\u2026": "...",   # ellipsis
        "\u00a0": " ",     # non-breaking space
        "\u2194": " <-> ", # left right arrow
        "\u2192": " -> ",  # right arrow
        "\u2190": " <- ",  # left arrow
        "\u2264": "<=",
        "\u2265": ">=",
        "\u2260": "!=",
    }

    def safe_text(text):
        """Safely encode text for PDF by replacing typographic Unicode symbols and non-printables."""
        if text is None:
            return ""
        s = str(text)
        for u_char, r_char in UNICODE_REPLACEMENTS.items():
            s = s.replace(u_char, r_char)
        # Normalize and strip unprintable characters
        s = unicodedata.normalize("NFKD", s)
        # Encode transliterated latin-1
        try:
            s = s.encode("latin-1", "replace").decode("latin-1")
        except Exception:
            s = s.encode("ascii", "replace").decode("ascii")
        return "".join(c for c in s if c.isprintable() or c in ("\n", "\t")).strip()

    is_dossier = len(headers) == 2 and headers[0] == "Property" and headers[1] == "Value"
    orientation = "P" if is_dossier else "L"

    # Font setup
    font_regular = "C:\\Windows\\Fonts\\arial.ttf" if _os.path.isfile("C:\\Windows\\Fonts\\arial.ttf") else None
    font_bold = "C:\\Windows\\Fonts\\arialbd.ttf" if _os.path.isfile("C:\\Windows\\Fonts\\arialbd.ttf") else None
    font_italic = "C:\\Windows\\Fonts\\ariali.ttf" if _os.path.isfile("C:\\Windows\\Fonts\\ariali.ttf") else None
    _font = "Arial" if font_regular else "helvetica"

    class ReportPDF(FPDF):
        def __init__(self, orientation="P"):
            super().__init__(orientation=orientation)
            if font_regular:
                self.add_font("Arial", "", font_regular)
                if font_bold:
                    self.add_font("Arial", "B", font_bold)
                if font_italic:
                    self.add_font("Arial", "I", font_italic)

        def header(self):
            self.set_font(_font, "B", 13)
            self.set_text_color(15, 23, 42)
            self.cell(0, 7, "SAKSHA Police Intelligence & Analytics Platform", align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_font(_font, "B", 10.5)
            self.set_text_color(30, 111, 217)
            self.cell(0, 6, safe_text(title), align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(203, 213, 225)
            self.line(10, self.get_y() + 2, self.w - 10, self.get_y() + 2)
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font(_font, "I", 8)
            self.set_text_color(100, 116, 139)
            self.cell(0, 10, f"SAKSHA Platform  |  Page {self.page_no()} of {{nb}}  |  CONFIDENTIAL LAW-ENFORCEMENT REPORT", align="C")

    pdf = ReportPDF(orientation=orientation)
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    filter_str = ", ".join(f"{k}={v}" for k, v in filters.items()) if filters else "None"

    # Metadata Box
    pdf.set_fill_color(248, 250, 252)
    pdf.set_draw_color(226, 232, 240)
    meta_y = pdf.get_y()
    meta_w = pdf.w - 20
    pdf.rect(10, meta_y, meta_w, 20, "DF")
    pdf.set_xy(14, meta_y + 2)
    
    pdf.set_font(_font, "B", 8.5)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 5, "INTELLIGENCE REPORT METADATA & CONTROL", new_x="LMARGIN", new_y="NEXT")
    pdf.set_x(14)
    pdf.set_font(_font, "", 8)
    pdf.set_text_color(71, 85, 105)
    pdf.cell(90, 4.5, f"Generated At: {generated}")
    pdf.cell(0, 4.5, safe_text(f"Classification / Watermark: {filter_str}"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_x(14)
    pdf.cell(0, 4.5, f"Total Data Attributes: {len(rows)}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_y(meta_y + 23)

    if headers and rows:
        if is_dossier:
            # 2-Column Key-Value Dossier Layout
            prop_w = 60
            val_w = meta_w - prop_w
            
            # Header Row
            pdf.set_fill_color(30, 41, 59)
            pdf.set_draw_color(51, 65, 85)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font(_font, "B", 8.5)
            pdf.cell(prop_w, 7, "  PROPERTY / ATTRIBUTE", border=1, fill=True)
            pdf.cell(val_w, 7, "  INTELLIGENCE VALUE & PARTICULARS", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

            for row_idx, row in enumerate(rows):
                prop_text = safe_text(row.get("Property", ""))
                val_text = safe_text(row.get("Value", ""))

                pdf.set_font(_font, "", 8)
                lines = pdf.multi_cell(val_w - 4, 4.5, val_text, dry_run=True, output="LINES")
                row_h = max(7, len(lines) * 4.5 + 3)

                if pdf.get_y() + row_h > pdf.h - 20:
                    pdf.add_page()
                    pdf.set_fill_color(30, 41, 59)
                    pdf.set_draw_color(51, 65, 85)
                    pdf.set_text_color(255, 255, 255)
                    pdf.set_font(_font, "B", 8.5)
                    pdf.cell(prop_w, 7, "  PROPERTY / ATTRIBUTE", border=1, fill=True)
                    pdf.cell(val_w, 7, "  INTELLIGENCE VALUE & PARTICULARS", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

                curr_x = 10
                curr_y = pdf.get_y()

                # Draw property cell
                pdf.set_fill_color(241, 245, 249) if row_idx % 2 == 0 else pdf.set_fill_color(248, 250, 252)
                pdf.set_draw_color(203, 213, 225)
                pdf.rect(curr_x, curr_y, prop_w, row_h, "DF")
                pdf.set_xy(curr_x + 2, curr_y + 1.5)
                pdf.set_font(_font, "B", 8)
                pdf.set_text_color(15, 23, 42)
                pdf.multi_cell(prop_w - 4, 4.2, prop_text)

                # Draw value cell
                pdf.set_fill_color(255, 255, 255) if row_idx % 2 == 0 else pdf.set_fill_color(250, 250, 250)
                pdf.rect(curr_x + prop_w, curr_y, val_w, row_h, "DF")
                pdf.set_xy(curr_x + prop_w + 2, curr_y + 1.5)
                pdf.set_font(_font, "", 8)
                pdf.set_text_color(30, 41, 59)
                pdf.multi_cell(val_w - 4, 4.5, val_text)

                pdf.set_xy(curr_x, curr_y + row_h)

        else:
            # Multi-column Tabular Report Layout
            col_w = max(22, (meta_w) / len(headers))
            pdf.set_fill_color(30, 41, 59)
            pdf.set_draw_color(51, 65, 85)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font(_font, "B", 8)
            for h in headers:
                pdf.cell(col_w, 7, safe_text(h.replace("_", " ").title())[:20], border=1, fill=True)
            pdf.ln()

            pdf.set_text_color(30, 41, 59)
            pdf.set_font(_font, "", 7.5)
            for row in rows:
                if pdf.get_y() + 7 > pdf.h - 20:
                    pdf.add_page()
                    pdf.set_fill_color(30, 41, 59)
                    pdf.set_draw_color(51, 65, 85)
                    pdf.set_text_color(255, 255, 255)
                    pdf.set_font(_font, "B", 8)
                    for h in headers:
                        pdf.cell(col_w, 7, safe_text(h.replace("_", " ").title())[:20], border=1, fill=True)
                    pdf.ln()
                    pdf.set_text_color(30, 41, 59)
                    pdf.set_font(_font, "", 7.5)

                for h in headers:
                    val = safe_text(row.get(h, ""))[:40]
                    pdf.cell(col_w, 6, val, border=1)
                pdf.ln()

    else:
        pdf.set_font(_font, "I", 9)
        pdf.cell(0, 8, "No records found matching the requested criteria.", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())


def _generate_docx(title: str, filters: dict, headers: list[str], rows: list[dict]) -> bytes:
    doc = Document()
    
    h0 = doc.add_heading("SAKSHA Police Intelligence & Analytics Platform", level=0)
    if h0.runs:
        h0.runs[0].font.bold = True
    
    doc.add_heading(title, level=1)
    
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    filter_str = ", ".join(f"{k}={v}" for k, v in filters.items()) if filters else "None"

    doc.add_heading("Report Metadata & Summary", level=2)
    doc.add_paragraph(f"Generated Date & Time: {generated}")
    doc.add_paragraph(f"Applied Filters: {filter_str}")
    doc.add_paragraph(f"Total Records: {len(rows)}")

    doc.add_heading("Report Details", level=2)

    if not headers or not rows:
        doc.add_paragraph("No data available for the requested criteria.")
    else:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h.replace("_", " ").title()
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            
        for row in rows:
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                row_cells[i].text = _clean_text(row.get(h))
                
    doc.add_paragraph()
    notice_p = doc.add_paragraph("SECURITY COMPLIANCE ACT NOTICE: CONFIDENTIAL LAW-ENFORCEMENT REPORT")
    if notice_p.runs:
        notice_p.runs[0].font.italic = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _create_report_record(db: Session, current_user: User, report_type: str, export_format: str, filters: dict[str, Any]) -> Report:
    report = Report(
        template=f"{report_type}_report",
        report_type=report_type,
        title=f"{report_type.title()} Report",
        requested_by_id=current_user.id,
        district=filters.get("district"),
        date_from=filters.get("date_from") or None,
        date_to=filters.get("date_to") or None,
        format=export_format,
        status=REPORT_STATUS_GENERATED,
        provenance=report_service.legacy_report_provenance(db, report_type, [], None),
        generation_method=GEN_METHOD_DATABASE_EXPORT,
        source_record_count=0,
        evidence_count=0,
        file_url=f"/api/v2/reports/{report_type}/export/{export_format}",
    )
    db.add(report)
    db.flush()
    return report


@router.get("", response_model=PaginatedResponse[ReportOut])
def list_reports(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    status: str | None = None,
    report_type: str | None = None,
    case_id: uuid.UUID | None = None,
    search: str | None = None,
    date_from: datetime | None = None,
    date_to: datetime | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Paginated, scoped report directory (§32/§34).

    Non-admin users only ever see reports they created; admins see everything.
    """
    query = db.query(Report)
    if current_user.role.name != ROLE_ADMIN:
        query = query.filter(Report.requested_by_id == current_user.id)
    if status:
        query = query.filter(Report.status == status)
    if report_type:
        query = query.filter(Report.report_type == report_type)
    if case_id:
        query = query.filter(Report.case_id == case_id)
    if search:
        query = query.filter(or_(
            Report.title.ilike(f"%{search}%"),
            Report.template.ilike(f"%{search}%"),
            Report.integrity_hash.ilike(f"%{search}%"),
            Report.provenance.ilike(f"%{search}%"),
        ))
    if date_from:
        query = query.filter(Report.created_at >= date_from)
    if date_to:
        query = query.filter(Report.created_at <= date_to)
    total = query.count()
    items = query.order_by(desc(Report.created_at)).offset((page - 1) * page_size).limit(page_size).all()
    return PaginatedResponse[ReportOut](total=total, page=page, page_size=page_size, results=items)


@router.get("/statistics/summary")
def report_statistics(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return {
        "cases": db.query(func.count(CrimeCase.id)).scalar() or 0,
        "officers": db.query(func.count(Officer.id)).scalar() or 0,
        "criminals": db.query(func.count(Criminal.id)).scalar() or 0,
        "evidence": db.query(func.count(Evidence.id)).scalar() or 0,
        "managed_reports": db.query(func.count(Report.id)).scalar() or 0,
    }


# --------------------------------------------------------------------------- #
# Issue #176 — production report lifecycle (draft -> … -> final -> archived)
# --------------------------------------------------------------------------- #
def _client_ip(request: Request) -> str | None:
    return request.client.host if request.client else None


def _load_report_or_404(db: Session, report_id: UUID) -> Report:
    report = db.query(Report).filter(Report.id == report_id).first()
    if report is None:
        raise NotFoundException("Report not found")
    return report


def _snapshot_rows_as_dicts(headers: list[str], rows: list) -> list[dict]:
    """Convert stored array-rows into the dict-row shape the exporters expect."""
    result = []
    for row in rows:
        if isinstance(row, dict):
            result.append(row)
            continue
        if isinstance(row, (list, tuple)):
            record = {}
            for i, h in enumerate(headers):
                record[h] = row[i] if i < len(row) else ""
            result.append(record)
        else:
            result.append({headers[0]: row} if headers else {})
    return result


def _render_snapshot_response(
    report: Report,
    snapshot: dict,
    export_format: str,
) -> Response:
    """Render a report from its stored snapshot (final reports are immutable).

    Never uses user-supplied filenames for filesystem access — the file bytes
    are produced in-memory and streamed back (§21/§22).
    """
    headers = list(snapshot.get("headers") or [])
    rows = _snapshot_rows_as_dicts(headers, list(snapshot.get("rows") or []))
    title = report.title or f"{report.report_type.title()} Report"
    filters = {
        "Report ID": str(report.id),
        "Report Type": report.report_type,
        "Version": f"v{report.version}",
        "Provenance": report.provenance,
        "Status": report.status,
        "Integrity Hash": report.integrity_hash or "pending",
    }
    filename = f"saksha_{report.report_type}_report_v{report.version}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}"

    if export_format == "csv":
        return _csv_response(filename, headers, rows)
    if export_format == "xlsx":
        return _xlsx_response(filename, title, filters, headers, rows)
    if export_format == "docx":
        docx_bytes = _generate_docx(title, filters, headers, rows)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
        )
    if export_format == "txt":
        txt_bytes = _generate_txt(title, filters, headers, rows)
        return Response(content=txt_bytes, media_type="text/plain; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{filename}.txt"'})
    pdf_bytes = _generate_pdf(title, filters, headers, rows)
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'})


@router.post("", response_model=ReportDetailOut)
def create_lifecycle_report(
    payload: ReportCreateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create a DRAFT report record (§4). The acting user is taken from the
    authenticated session — never trusted from the client (§19)."""
    report = report_service.create_report(db, current_user, payload.model_dump())
    db.commit()
    db.refresh(report)
    return report_service.serialize_report_details(db, report)


@router.get("/{report_id:uuid}", response_model=ReportDetailOut)
def get_report_detail(
    report_id: UUID,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    audit_service.log_action(
        db, current_user, "REPORT_VIEW", "Report", str(report.id),
        ip_address=_client_ip(request),
        metadata_json=json.dumps({"status": report.status, "version": report.version}),
    )
    return report_service.serialize_report_details(db, report)


@router.post("/{report_id:uuid}/validate", response_model=dict)
def validate_report_references(
    report_id: UUID,
    payload: ReportValidateRequest,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Validate that referenced source/evidence records exist (§9)."""
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    result = report_service.validate_references(
        db, [s.model_dump() for s in payload.sources], payload.evidence_ids
    )
    audit_service.log_action(
        db, current_user, "REPORT_AI_VALIDATION", "Report", str(report.id),
        details=f"verified={len(result['verified_records'])} missing={len(result['missing_records'])}",
        ip_address=_client_ip(request),
        result="success" if result["can_finalize_as_verified"] else "failure",
    )
    return result


@router.post("/{report_id:uuid}/generate", response_model=ReportDetailOut)
def generate_lifecycle_report(
    report_id: UUID,
    payload: ReportGeneratePayload,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    try:
        report = report_service.generate_report(
            db, current_user, report, payload.model_dump(), ip_address=_client_ip(request)
        )
        db.commit()
    except Exception:
        # Persist the FAILED state so a broken run never looks successful (§25).
        db.commit()
        raise
    db.refresh(report)
    return report_service.serialize_report_details(db, report)


@router.get("/{report_id:uuid}/versions", response_model=list)
def list_report_versions(
    report_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    return report_service.serialize_report_details(db, report).get("versions", [])


@router.post("/{report_id:uuid}/versions", response_model=ReportDetailOut)
def create_report_version(
    report_id: UUID,
    payload: ReportVersionCreateRequest,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Explicit versioning (§10/§11): never silently overwrite a previous
    version; every change is an auditable new version with a reason."""
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    if report.status == report_service.REPORT_STATUS_ARCHIVED:
        raise ConflictException("Archived reports cannot be versioned")
    report_service.create_version(
        db, current_user, report,
        reason=payload.reason, new_content=payload.content, ip_address=_client_ip(request),
    )
    db.commit()
    db.refresh(report)
    return report_service.serialize_report_details(db, report)


@router.post("/{report_id:uuid}/review", response_model=ReportDetailOut)
def review_lifecycle_report(
    report_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    request: Request = None,
    payload: ReportReviewRequest | None = None,
):
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    notes = payload.notes if payload else None
    report = report_service.start_review(
        db, current_user, report, notes=notes, ip_address=_client_ip(request)
    )
    db.commit()
    db.refresh(report)
    return report_service.serialize_report_details(db, report)


@router.post("/{report_id:uuid}/finalize", response_model=ReportDetailOut)
def finalize_lifecycle_report(
    report_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    request: Request = None,
    payload: ReportReviewRequest | None = None,
):
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    notes = payload.notes if payload else None
    report = report_service.finalize_report(
        db, current_user, report, notes=notes, ip_address=_client_ip(request)
    )
    db.commit()
    db.refresh(report)
    return report_service.serialize_report_details(db, report)


@router.post("/{report_id:uuid}/archive", response_model=ReportDetailOut)
def archive_lifecycle_report(
    report_id: UUID,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    report = report_service.archive_report(db, current_user, report, ip_address=_client_ip(request))
    db.commit()
    db.refresh(report)
    return report_service.serialize_report_details(db, report)


@router.get("/{report_id:uuid}/download")
def download_lifecycle_report(
    report_id: UUID,
    request: Request,
    export_format: str = Query("pdf", pattern="^(pdf|csv|docx|txt|xlsx)$"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Authorized, audited download of a managed report (§21).

    Final reports render from their immutable stored snapshot — their content
    cannot silently change when the source DB changes (§27/§28).
    """
    report = _load_report_or_404(db, report_id)
    report_service.require_report_access(current_user, report)
    snapshot = report_service._load_snapshot(report)
    if not snapshot.get("headers") and not snapshot.get("rows"):
        raise ConflictException("Report has no generated content to download")
    audit_service.log_action(
        db, current_user, "REPORT_DOWNLOAD", "Report", str(report.id),
        details=f"format={export_format}; version=v{report.version}",
        ip_address=_client_ip(request),
        metadata_json=json.dumps({"format": export_format, "version": report.version}),
    )
    db.commit()
    return _render_snapshot_response(report, snapshot, export_format)


@router.get("/{report_id:uuid}/audit", response_model=dict)
def report_audit_history(
    report_id: UUID,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles(ROLE_ADMIN)),
):
    """Administrative audit history for a report (§13/§17/§33).

    Only admins may query audit logs; ordinary users are denied.
    """
    _load_report_or_404(db, report_id)
    query = db.query(AuditLog).filter(
        AuditLog.resource_type == "Report",
        AuditLog.resource_id == str(report_id),
    )
    total = query.count()
    items = query.order_by(desc(AuditLog.timestamp)).offset((page - 1) * page_size).limit(page_size).all()
    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "results": [report_service.serialize_audit_entry(e) for e in items],
    }


@router.get("/{report_type}")
def preview_report(
    report_type: str,
    search: str | None = None,
    status: str | None = None,
    district: str | None = None,
    date_from: datetime | None = None,
    date_to: datetime | None = None,
    sort_by: str = "created_at",
    sort_order: str = Query("desc", pattern="^(asc|desc)$"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if report_type not in REPORT_TYPES:
        return Response(status_code=404, content="Unknown report type")
    query, headers, mapper = _report_query(db, report_type, search, status, district, date_from, date_to, sort_by, sort_order)
    total = query.count()
    items = query.offset((page - 1) * page_size).limit(page_size).all()
    return {
        "report_type": report_type,
        "headers": headers,
        "filters": _filters_dict(search=search, status=status, district=district, date_from=date_from, date_to=date_to, sort_by=sort_by, sort_order=sort_order),
        "total": total,
        "page": page,
        "page_size": page_size,
        "results": [mapper(item) for item in items],
    }


@router.post("/{report_type}/generate")
def generate_report(
    report_type: str,
    request: Request,
    export_format: str = Query("pdf", pattern="^(pdf|csv|docx|txt|xlsx)$"),
    search: str | None = None,
    status: str | None = None,
    district: str | None = None,
    date_from: datetime | None = None,
    date_to: datetime | None = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if report_type not in REPORT_TYPES:
        return Response(status_code=404, content="Unknown report type")
    filters = _filters_dict(search=search, status=status, district=district, date_from=date_from, date_to=date_to)
    report = _create_report_record(db, current_user, report_type, export_format, {"district": district, "date_from": date_from, "date_to": date_to})
    audit_service.log_action(
        db, current_user, "REPORT_GENERATE", "Report", str(report.id),
        details=str(filters), ip_address=request.client.host if request.client else None,
        metadata_json=json.dumps({"type": report_type, "provenance": report.provenance}),
    )
    db.commit()
    db.refresh(report)
    return {"id": str(report.id), "status": report.status, "format": report.format, "file_url": report.file_url}


@router.get("/{report_type}/export/{export_format}")
def export_report(
    report_type: str,
    export_format: str,
    request: Request,
    search: str | None = None,
    status: str | None = None,
    district: str | None = None,
    date_from: datetime | None = None,
    date_to: datetime | None = None,
    sort_by: str = "created_at",
    sort_order: str = Query("desc", pattern="^(asc|desc)$"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if report_type not in REPORT_TYPES or export_format not in EXPORT_FORMATS:
        return Response(status_code=404, content="Unknown export")
    query, headers, mapper = _report_query(db, report_type, search, status, district, date_from, date_to, sort_by, sort_order)
    rows = [mapper(item) for item in query.limit(5000).all()]
    filters = _filters_dict(search=search, status=status, district=district, date_from=date_from, date_to=date_to, sort_by=sort_by, sort_order=sort_order)
    report = _create_report_record(db, current_user, report_type, export_format, {"district": district, "date_from": date_from, "date_to": date_to})
    report.source_record_count = len(rows)
    if report_type == "evidence":
        report.evidence_count = len(rows)
    report.provenance = report_service.legacy_report_provenance(db, report_type, rows, report)
    audit_service.log_action(
        db, current_user, "REPORT_EXPORT", "Report", str(report.id),
        details=f"{export_format}:{filters}", ip_address=request.client.host if request.client else None,
        metadata_json=json.dumps({
            "type": report_type, "provenance": report.provenance,
            "record_count": len(rows), "format": export_format,
        }),
    )
    db.commit()
    filename = f"saksha_{report_type}_report_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}"
    title = f"{report_type.title()} Report"
    
    if export_format == "csv":
        return _csv_response(filename, headers, rows)
    elif export_format == "xlsx":
        return _xlsx_response(filename, title, filters, headers, rows)
    elif export_format == "docx":
        docx_bytes = _generate_docx(title, filters, headers, rows)
        return Response(content=docx_bytes, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'})
    elif export_format == "txt":
        txt_bytes = _generate_txt(title, filters, headers, rows)
        return Response(content=txt_bytes, media_type="text/plain; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{filename}.txt"'})
    
    pdf_bytes = _generate_pdf(title, filters, headers, rows)
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'})


class DossierPayload(BaseModel):
    title: str
    data: dict[str, Any]
    watermark: str = ""


IGNORED_INTERNAL_KEYS = {
    "x", "y", "z", "vx", "vy", "vz", "index", "__threeObj", "__line", 
    "__indexColor", "__nodeColor", "__isHovered", "isHovered", "selected",
    "source", "target", "fx", "fy", "fz", "neo4j_node_id", "extra", "color"
}

KEY_TITLE_MAPPINGS = {
    "id": "Profile Reference ID",
    "name": "Subject Full Name",
    "category": "Classification",
    "riskScore": "Threat & Risk Assessment",
    "risk_score": "Threat & Risk Assessment",
    "details": "Modus Operandi & Narrative",
    "casesCount": "Active FIR Cases",
    "cases_count": "Active FIR Cases",
    "phone": "Contact Telephone",
    "gangAffiliation": "Gang / Syndicate Affiliation",
    "gang_affiliation": "Gang / Syndicate Affiliation",
    "status": "Operational Status",
    "district": "Jurisdiction District",
    "date": "Incident / Record Date",
    "isSeed": "Intelligence Grounding Scope",
    "is_seed": "Intelligence Grounding Scope",
    "suspectName": "Subject Name",
    "casesFilings": "Recorded FIR Filings",
    "links": "Network Associations",
    "activeSuspects": "Identified Suspect Entities",
    "relationType": "Analysis Classification",
    "totalNodes": "Total Network Nodes",
    "totalEdges": "Total Relationship Edges",
    "provenanceSummary": "Provenance Breakdown",
    "relationEdges": "Verified & Analytical Linkages",
}


@router.post("/dossier/export/{export_format}")
def export_dossier(
    export_format: str,
    payload: DossierPayload,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if export_format not in EXPORT_FORMATS:
        return Response(status_code=404, content="Unknown export format")
        
    audit_service.log_action(db, current_user, "DOSSIER_EXPORT", "Dossier", payload.title, details=export_format)
    
    headers = ["Property", "Value"]
    rows = []
    for k, v in payload.data.items():
        if k in IGNORED_INTERNAL_KEYS:
            continue
        if v is None or (isinstance(v, str) and not v.strip()):
            continue
        if isinstance(v, (list, dict)) and len(v) == 0:
            continue

        prop_name = KEY_TITLE_MAPPINGS.get(k, str(k).replace("_", " ").title())

        if k in ("riskScore", "risk_score") and isinstance(v, (int, float)):
            val_str = f"{v} / 100 ({'Critical Threat' if v >= 75 else 'High Threat' if v >= 50 else 'Moderate Risk'})"
        elif k in ("category", "status") and isinstance(v, str):
            val_str = v.replace("_", " ").upper()
        elif k in ("isSeed", "is_seed") and isinstance(v, bool):
            val_str = "Demonstration / Training Fixture" if v else "Live Police Operational Intelligence"
        elif k in ("casesCount", "cases_count") and isinstance(v, (int, float)):
            val_str = f"{int(v)} Registered FIR{'s' if v != 1 else ''}"
        else:
            val_str = _format_human_readable_value(v)

        if val_str:
            rows.append({"Property": prop_name, "Value": val_str})

    filters = {"Watermark": payload.watermark} if payload.watermark else {}
    filename = f"ksp_{payload.title.lower().replace(' ', '_')}"
    
    if export_format == "csv":
        return _csv_response(filename, headers, rows)
    elif export_format == "xlsx":
        return _xlsx_response(filename, payload.title, filters, headers, rows)
    elif export_format == "docx":
        docx_bytes = _generate_docx(payload.title, filters, headers, rows)
        return Response(content=docx_bytes, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'})
    elif export_format == "txt":
        txt_bytes = _generate_txt(payload.title, filters, headers, rows)
        return Response(content=txt_bytes, media_type="text/plain; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{filename}.txt"'})
    
    pdf_bytes = _generate_pdf(payload.title, filters, headers, rows)
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'})

